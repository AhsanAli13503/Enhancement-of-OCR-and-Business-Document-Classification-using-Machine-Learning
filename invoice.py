
import cv2    
import numpy as np
import glob
import os
from docx import Document
import pytesseract
from charsegment import charsegment
from imgtotext import img_to_text
typeOfimg = "leinre"

def imageReader(path):
    images = []
    for i in range(path):
        images.append(cv2.imread("tablesegments\\"+str(i)+".png",0))
    return images

def get_contour_precedence(contour, cols):
    tolerance_factor = 10
    origin = cv2.boundingRect(contour)
    return ((origin[1] // tolerance_factor) * tolerance_factor) * cols + origin[0]

def preprocessing(img):
    
    image = cv2.equalizeHist(img)
    filtered = cv2.adaptiveThreshold(image.astype(np.uint8), 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY, 9, 41)
    kernel = np.ones((1, 1), np.uint8)
    opening = cv2.morphologyEx(filtered, cv2.MORPH_OPEN, kernel)
    closing = cv2.morphologyEx(opening, cv2.MORPH_CLOSE, kernel)
    or_image = cv2.bitwise_or(image, closing)
    ret,thresh1 = cv2.threshold(or_image,0, 255,cv2.THRESH_OTSU|cv2.THRESH_BINARY_INV)
    
    rect_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (20, 3))
    dilation = cv2.dilate(thresh1, rect_kernel, iterations = 2)

    return dilation,image
   
def table_processing(img):
    files = glob.glob('tablesegments\\*')
    for f in files:
            os.remove(f)
    img1 = cv2.equalizeHist(img) 
# Thresholding the image
    (thresh, img_bin) = cv2.threshold(img1, 128, 255,cv2.THRESH_BINARY|     cv2.THRESH_OTSU)
# Invert the image
    img_bin = 255-img_bin 
    
# Defining a kernel length
    kernel_length = np.array(img1).shape[1]//80
 
# A verticle kernel of (1 X kernel_length), which will detect all the verticle lines from the image.
    verticle_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, kernel_length))
# A horizontal kernel of (kernel_length X 1), which will help to detect all the horizontal line from the image.
    hori_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (kernel_length, 1))
# A kernel of (3 X 3) ones.
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))

# Morphological operation to detect vertical lines from an image
    img_temp1 = cv2.erode(img_bin, verticle_kernel, iterations=7)
    verticle_lines_img = cv2.dilate(img_temp1, verticle_kernel, iterations=7)
    
# Morphological operation to detect horizontal lines from an image
    img_temp2 = cv2.erode(img_bin, hori_kernel, iterations=3)
    horizontal_lines_img = cv2.dilate(img_temp2, hori_kernel, iterations=3)
    
# Weighting parameters, this will decide the quantity of an image to be added to make a new image.
    alpha = 0.5
    beta = 1.0 - alpha
# This function helps to add two image with specific weight parameter to get a third image as summation of two image.
    img_final_bin = cv2.addWeighted(verticle_lines_img, alpha, horizontal_lines_img, beta, 0.0)
    img_final_bin = cv2.erode(~img_final_bin, kernel, iterations=1)
    (thresh, img_final_bin) = cv2.threshold(img_final_bin, 128,255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
    
# Find contours for image, which will detect all the boxes
    im2, contours, hierarchy = cv2.findContours(img_final_bin, cv2.RETR_EXTERNAL , cv2.CHAIN_APPROX_SIMPLE)
# Sort all the contours by top to bottom.
    contours.sort(key=lambda x:get_contour_precedence(x, img_final_bin.shape[1])) 
    tablecontours = 0
    for c in contours:
        # Returns the location and width,height for every contour
        x, y, w, h = cv2.boundingRect(c)
        new_img = img[y:y+h, x:x+w]
        
        if h<150:
            cv2.imwrite("tablesegments\\"+str(tablecontours) + '.png', new_img)
            tablecontours += 1
    im2, contours, hierarchy = cv2.findContours(verticle_lines_img, cv2.RETR_EXTERNAL , cv2.CHAIN_APPROX_SIMPLE)     
    return tablecontours,len(contours)-1
def invoice(img,tableimg):
    document = Document()    
    
    dilation,enhanceimage= preprocessing(img)
    
    _ ,contours, _ = cv2.findContours(dilation, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
    contours.sort(key=lambda x:get_contour_precedence(x, img.shape[1]))
    
    areaThr = 61000
    i = 0
    
    number_of_tables=0
    data1=[]
    ypre = 0
    xpre = 0
    datapre=""
    p=""
    for cnt in contours:
            x, y, width, height = cv2.boundingRect(cnt)
            area = cv2.contourArea(cnt)
            if (area > areaThr):
                number_of_tables=number_of_tables+1            
                i = i + 1
                table = img[y:y+height-1, x:x+width-1]
                """
                cv2.imshow("Table",table)
                cv2.waitKey()
                cv2.destroyAllWindows()
                """
                data,vertical=table_processing(table)
                cols= vertical
                rows1 = int(data/vertical)
                data_images=imageReader(data)
                table = document.add_table(0, cols)
                itera=0
                for i in range(rows1):
                    tableimg=1
                    row_cells = table.add_row().cells
                    for j in range(cols):
                        roi = cv2.resize(data_images[itera],None,fx=4, fy=4, interpolation = cv2.INTER_CUBIC)
                        text = pytesseract.image_to_string(roi)
                        #if text == "":
                            #segment=charsegment(roi)
                            #text = img_to_text(segment)
                        row_cells[j].text = text
                        data1.append(text)
                        itera+=1
            else:
               
                table = img[y:y+height-1, x:x+width-1]
                roi = cv2.resize(table,None,fx=4, fy=4, interpolation = cv2.INTER_CUBIC)
                text = pytesseract.image_to_string(roi)
                
               # if text == "":
                    #segment=charsegment(roi)
                    #text = " "+img_to_text(segment)
                data1.append(text)
                if ypre!=0 and y==ypre:
                    space = x-xpre
                    space = int(space/7)
                    datapace=""
                    for i in range(space):
                        datapace=datapace+" "
                    text = datapace+text
                    p.add_run(text)
                else:
                    if x<100:
                        if len(datapre)==0:
                            p=document.add_paragraph(text)
                            
                        else:
                            if(len(datapre))>70:
                                p.add_run(text)
                                
                            else:
                                p=document.add_paragraph(text)
                                p.alignment = 0
                                
                    elif x>=100 and x<=300:
                        p = document.add_paragraph(text)
                        p.alignment = 1
                    
                    else:
                        p = document.add_paragraph(text)
                        p.alignment = 2
            ypre=  y
            xpre = x
            datapre=text
    
    document.save("static\\document.docx")
    print (tableimg)
    return data1,typeOfimg,tableimg
img = cv2.imread("Capture.png",0)
invoice(img,0)

