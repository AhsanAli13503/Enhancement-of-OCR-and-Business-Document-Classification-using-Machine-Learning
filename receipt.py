import cv2
import numpy as np
import os
import pytesseract
from docx import Document
import imutils
from PIL import Image
typeOfimg = "pos"

def receipt(image,tableimg):
    doc = Document()
    orig= image.copy()
    process_height = 500
    ratio = image.shape[0] / float(process_height)
    gray = image
    gray = cv2.GaussianBlur(gray, (5, 5), 0)
    edged = cv2.Canny(gray, 10, 50)
    cnts = cv2.findContours(edged.copy(), cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
    cnts = cnts[0] if imutils.is_cv2() else cnts[1]
    cnts = sorted(cnts, key=cv2.contourArea, reverse=True)[:5]
    screenCnt = None
    
    for c in cnts:
        peri = cv2.arcLength(c, True)
        approx = cv2.approxPolyDP(c, 0.02 * peri, True)
        if len(approx) == 4:
            screenCnt = approx
            break
    x, y, w, h = cv2.boundingRect(screenCnt)
    if x==0:
        img = cv2.resize(orig, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
        img = cv2.GaussianBlur(img, (5, 5), 0)
        img = cv2.medianBlur(img, 3)
        img = cv2.bilateralFilter(img,9,75,75)
        text = pytesseract.image_to_string(img)
        p=doc.add_paragraph(text)
        doc.save("static\\document.docx")
        return text,typeOfimg,tableimg
    else:
        newimg =orig[y :y + h, x:x + w]
        img = cv2.resize(newimg, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
        img = cv2.GaussianBlur(img, (5, 5), 0)
        img = cv2.medianBlur(img, 3)
        img = cv2.bilateralFilter(img,9,75,75)
        text = pytesseract.image_to_string(img)
        p=doc.add_paragraph(text)
        doc.save("static\\document.docx")
        return text,typeOfimg,tableimg
   