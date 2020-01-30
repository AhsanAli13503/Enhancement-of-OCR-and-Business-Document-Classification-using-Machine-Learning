import cv2
from docx import Document
import pytesseract
from docx import Document
from charsegment import charsegment
from imgtotext import img_to_text


def preprocess(image):
    img = cv2.GaussianBlur(image, (5, 5), 0)
    __,bimg=cv2.threshold(img, 0, 255,cv2.THRESH_OTSU|cv2.THRESH_BINARY_INV)
    return bimg
def get_contour_precedence(contour, cols):
    tolerance_factor = 10
    origin = cv2.boundingRect(contour)
    return ((origin[1] // tolerance_factor) * tolerance_factor) * cols + origin[0]

def findandSortExternalContours(bimg,orignal): 
    rect_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (20, 3))
    dilation = cv2.dilate(bimg, rect_kernel, iterations = 1)
    __,contours, __ = cv2.findContours(dilation, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
    contours.sort(key=lambda x:get_contour_precedence(x, bimg.shape[1]))
    segments = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        roi =orignal[y :y + h, x:x + w] 
        segments.append(roi)
    return segments

def extract_and_text(img):
    text = pytesseract.image_to_string(img)
    if len(text) == 0:
        segment=charsegment(img)
        text = img_to_text(segment)
    if len(text) == 1:
        return ""
    return text
def letter(orignal):
    doc = Document()
    bimg = preprocess(orignal)
    exseg=findandSortExternalContours(bimg,orignal)
    data1 = []
    for data in exseg:
        text = extract_and_text(data)
        data1.append(text)
    p=doc.add_paragraph()
    for i in range(len(data1)):
        if i == 0:
            p=doc.add_paragraph(data1[i])
        else:
            if len(data1[i])>0:
                if len(data1[i-1])>80:
                    p.add_run(data1[i])
                else:
                     p=doc.add_paragraph(data1[i])
    doc.save("static\\document.docx")
    return data1
    
        